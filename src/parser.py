"""Document parsing for various file formats."""

import hashlib
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import tiktoken


class DocumentParser:
    """Parse documents of various formats."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
    
    def parse_file(self, file_path: Path) -> Dict[str, Any]:
        """Parse a file and extract its content and metadata."""
        import logging
        logger = logging.getLogger(__name__)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_stats = file_path.stat()
        file_hash = self._calculate_file_hash(file_path)
        file_type = file_path.suffix.lower()
        
        logger.info(f"    → Extracting text from {file_type} file...")
        
        metadata = {
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_size": file_stats.st_size,
            "modified_time": datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            "file_hash": file_hash,
            "file_type": file_type
        }
        
        # Extract text based on file type
        text = self._extract_text(file_path)
        logger.info(f"    → Extracted {len(text):,} characters")
        
        logger.info(f"    → Creating text chunks (size: {self.chunk_size}, overlap: {self.chunk_overlap})...")
        # Create chunks
        chunks = self._create_chunks(text)
        
        return {
            "metadata": metadata,
            "text": text,
            "chunks": chunks,
            "num_chunks": len(chunks),
            "total_chars": len(text),
            "total_tokens": len(self.tokenizer.encode(text))
        }
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA256 hash of file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from various file formats."""
        file_ext = file_path.suffix.lower()
        
        if file_ext == ".pdf":
            return self._extract_pdf_text(file_path)
        elif file_ext in [".docx", ".doc"]:
            return self._extract_docx_text(file_path)
        elif file_ext in [".txt", ".md", ".rtf"]:
            return self._extract_plain_text(file_path)
        else:
            # Try to read as plain text
            try:
                return self._extract_plain_text(file_path)
            except Exception:
                raise ValueError(f"Unsupported file type: {file_ext}")
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text_parts = []
        try:
            with fitz.open(str(file_path)) as pdf:
                for page_num, page in enumerate(pdf, 1):
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(f"[Page {page_num}]\n{text}")
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {e}")
        
        return "\n\n".join(text_parts)
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from Word document."""
        try:
            doc = DocxDocument(str(file_path))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise ValueError(f"Error parsing Word document: {e}")
    
    def _extract_plain_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with errors='ignore'
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"Error reading text file: {e}")
    
    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create overlapping chunks from text."""
        import logging
        logger = logging.getLogger(__name__)
        
        if not text:
            return []
        
        text_length = len(text)
        estimated_chunks = (text_length // self.chunk_size) + 1
        logger.info(f"    → Estimated {estimated_chunks} chunks for {text_length:,} characters")
        
        chunks = []
        start = 0
        chunk_id = 0
        last_progress = 0
        
        while start < text_length:
            # Progress logging every 10% or every 100 chunks
            progress = int((start / text_length) * 100)
            if progress >= last_progress + 10 or chunk_id % 100 == 0:
                logger.info(f"    → Chunking progress: {progress}% ({chunk_id} chunks created)")
                last_progress = progress
            
            # Find the end of the chunk
            end = min(start + self.chunk_size, text_length)
            
            # Try to break at sentence or paragraph boundary
            if end < text_length:
                # Look for paragraph break first (better boundary)
                para_break = text.rfind('\n\n', start, end)
                if para_break > start:
                    end = para_break
                else:
                    # Look for sentence break
                    sentence_breaks = ['. ', '! ', '? ', '\n']
                    for break_char in sentence_breaks:
                        break_pos = text.rfind(break_char, start, end)
                        if break_pos > start:
                            end = break_pos + len(break_char)
                            break
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                # Only tokenize if chunk is not empty - this is the expensive operation
                token_count = len(self.tokenizer.encode(chunk_text))
                
                chunks.append({
                    "chunk_id": chunk_id,
                    "text": chunk_text,
                    "start_pos": start,
                    "end_pos": end,
                    "char_count": len(chunk_text),
                    "token_count": token_count
                })
                chunk_id += 1
            
            # Move to next chunk with overlap, but ensure we make progress
            new_start = max(end - self.chunk_overlap, start + 1)
            if new_start <= start:
                # Force progress to prevent infinite loop
                new_start = start + max(1, self.chunk_size // 2)
            
            start = new_start
            
            # Safety check to prevent infinite loops
            if chunk_id > 10000:  # Reasonable limit
                logger.warning(f"    → Chunk limit reached ({chunk_id}), stopping to prevent infinite loop")
                break
        
        logger.info(f"    → Created {len(chunks)} chunks")
        return chunks


def is_supported_file(file_path: Path, extensions: List[str]) -> bool:
    """Check if file is supported based on extension."""
    return file_path.suffix.lower() in extensions


def get_file_info(file_path: Path) -> Dict[str, Any]:
    """Get basic file information without parsing content."""
    if not file_path.exists():
        return None
    
    stats = file_path.stat()
    mime_type, _ = mimetypes.guess_type(str(file_path))
    
    return {
        "path": str(file_path),
        "name": file_path.name,
        "size": stats.st_size,
        "size_mb": round(stats.st_size / (1024 * 1024), 2),
        "modified": datetime.fromtimestamp(stats.st_mtime).isoformat(),
        "created": datetime.fromtimestamp(stats.st_ctime).isoformat(),
        "extension": file_path.suffix.lower(),
        "mime_type": mime_type
    }